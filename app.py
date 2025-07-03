import streamlit as st
import openai
import base64
import json
import requests
from ast import literal_eval
from io import BytesIO
from docx import Document

# Setup API keys from secrets
openai.api_key = st.secrets["OPENAI_API_KEY"]
AVETI_API_TOKEN = st.secrets["AVETI_API_TOKEN"]

# API configuration
api_url = "https://production.mobile.avetilearning.com/service/cms/api/v1/exercise/75020/questions"
headers = {
    "Authorization": f"Bearer {AVETI_API_TOKEN}",
    "Content-Type": "application/json"
}

st.title("🧠 MCQ Extractor (Perseus Style) from Multiple Images")
st.write("Upload multiple images with MCQs. Get results in structured JSON format like Perseus uses.")
debug = st.sidebar.checkbox("🔍 Show raw JSON from OpenAI")

uploaded_files = st.file_uploader(
    "Upload image(s) with MCQs", 
    type=["jpg", "jpeg", "png"], 
    accept_multiple_files=True
)

def extract_json_mcqs_from_image(image_bytes):
    base64_image = base64.b64encode(image_bytes).decode("utf-8")
    prompt = (
        "Extract **ALL** multiple choice questions (MCQs) from this image, even if they are partially visible or unclear. "
        "Do not skip any question. For each question, if an option or answer is unreadable, include a placeholder such as 'Unclear'. "
        "Respond with ONLY valid JSON in this format (one list of objects):\n\n"
        "[{\"question\": \"...\", \"options\": [\"...\", \"...\", \"...\", \"...\"], \"answer_index\": 1}]\n\n"
        "Do NOT include markdown, explanation, or any extra text. No ```json blocks. Only plain JSON."
    )
    response = openai.chat.completions.create(
        model="gpt-4o",
        temperature=0,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                ]
            }
        ],
        max_tokens=2000
    )
    return response.choices[0].message.content

def send_mcq_to_api(mcq):
    try:
        payload = {
            "question": {
                "content": mcq["question"] + "\n\n[[radio 1]]",
                "images": {},
                "widgets": {
                    "radio 1": {
                        "type": "radio",
                        "alignment": "default",
                        "static": False,
                        "graded": True,
                        "options": {
                            "choices": [
                                {"content": opt, "correct": (i == mcq["answer_index"])}
                                for i, opt in enumerate(mcq["options"])
                            ],
                            "randomize": True,
                            "multipleSelect": False,
                            "displayCount": None,
                            "hasNoneOfTheAbove": False,
                            "onePerLine": True,
                            "deselectEnabled": False
                        },
                        "version": {"major": 1, "minor": 0}
                    }
                }
            },
            "answerArea": {
                "calculator": False,
                "chi2Table": False,
                "periodicTable": False,
                "tTable": False,
                "zTable": False
            },
            "itemDataVersion": {
                "major": 0,
                "minor": 1
            },
            "hints": [
                {
                    "replace": False,
                    "content": "Hint 1: Think carefully.",
                    "images": {},
                    "widgets": {}
                },
                {
                    "replace": False,
                    "content": "The correct answer is:\n\n[[radio 1]]",
                    "images": {},
                    "widgets": {
                        "radio 1": {
                            "type": "radio",
                            "alignment": "default",
                            "static": True,
                            "graded": True,
                            "options": {
                                "choices": [
                                    {"content": opt, "correct": (i == mcq["answer_index"])}
                                    for i, opt in enumerate(mcq["options"])
                                ],
                                "randomize": False,
                                "multipleSelect": False,
                                "displayCount": None,
                                "hasNoneOfTheAbove": False,
                                "onePerLine": True,
                                "deselectEnabled": False
                            },
                            "version": {"major": 1, "minor": 0}
                        }
                    }
                }
            ]
        }

        # ✅ FIX: Wrap in `question_json` as required by the API
        response = requests.post(api_url, headers=headers, json={"question_json": payload})
        return response.status_code, response.text, payload
    except Exception as e:
        return 500, str(e), {}

if uploaded_files:
    all_mcqs = []
    raw_outputs = []

    for img in uploaded_files:
        st.image(img, caption=img.name, use_container_width=True)
        with st.spinner(f"Extracting MCQs from {img.name}..."):
            raw_json = extract_json_mcqs_from_image(img.read())
            raw_outputs.append((img.name, raw_json))

            if debug:
                st.text_area(f"🧾 Raw Output from {img.name}", raw_json, height=150)

            if raw_json.strip().startswith("```json"):
                raw_json = raw_json.strip().removeprefix("```json").removesuffix("```").strip()

            try:
                mcqs = json.loads(raw_json)
            except json.JSONDecodeError:
                st.warning(f"⚠️ {img.name} returned invalid JSON. Trying fallback parser...")
                try:
                    mcqs = literal_eval(raw_json)
                except Exception as e:
                    mcqs = []
                    st.error(f"❌ Could not parse fallback for {img.name}: {e}")

            if mcqs:
                for i, mcq in enumerate(mcqs, start=1):
                    status, resp, sent_payload = send_mcq_to_api(mcq)
                    if status in [200, 201]:
                        st.success(f"✅ Question {i} sent successfully.")
                        st.markdown("**Payload sent:**")
                        st.code(json.dumps(sent_payload, indent=2), language="json")
                        st.markdown(f"**Server response:** `{resp}`")
                    else:
                        st.error(f"❌ Failed to send Question {i} (Status Code: {status})")
                        st.markdown("**Payload attempted:**")
                        st.code(json.dumps(sent_payload, indent=2), language="json")
                        try:
                            error_json = json.loads(resp)
                            st.markdown("**Error details from API:**")
                            st.code(json.dumps(error_json, indent=2), language="json")
                        except:
                            st.markdown("**Raw error response:**")
                            st.code(resp, language="text")
                        print(f"[ERROR] Status {status} | Response: {resp}")
                all_mcqs.extend(mcqs)
            else:
                st.warning(f"⚠️ No valid MCQs could be parsed from {img.name}")

    st.subheader("🧾 Final Output (Combined from All Images)")
    if all_mcqs:
        formatted = json.dumps(all_mcqs, indent=2)
        st.text_area("📋 Structured Perseus JSON", formatted, height=400)
        st.download_button("📄 Download JSON", data=formatted, file_name="perseus_mcqs.json", mime="application/json")

        doc = Document()
        doc.add_heading("MCQs Extracted in Perseus Format", 0)
        for q in all_mcqs:
            doc.add_paragraph(f"Q: {q['question']}")
            for i, opt in enumerate(q["options"]):
                prefix = " ✅" if i == q["answer_index"] else ""
                doc.add_paragraph(f"  - {opt}{prefix}")
            doc.add_paragraph("")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("📄 Download as Word Document", data=buffer, file_name="mcqs_perseus.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("❗ No structured MCQs could be parsed. Check raw output below.")

    st.subheader("🧾 Raw Outputs (Unparsed)")
    for filename, raw in raw_outputs:
        st.text_area(f"🖼 Raw output from {filename}", raw, height=200)
